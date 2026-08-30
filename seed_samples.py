import os
import sys
import asyncio
import aiosqlite
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from PIL import Image, ImageDraw, ImageFont

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOADS_DIR = os.path.join(BASE_DIR, "uploads")
EBOOKS_DIR = os.path.join(UPLOADS_DIR, "ebooks")
COVERS_DIR = os.path.join(UPLOADS_DIR, "covers")
SAMPLES_DIR = os.path.join(UPLOADS_DIR, "samples")
DB_PATH = os.path.join(BASE_DIR, "store.db")

os.makedirs(EBOOKS_DIR, exist_ok=True)
os.makedirs(COVERS_DIR, exist_ok=True)
os.makedirs(SAMPLES_DIR, exist_ok=True)

def generate_cover_image(title: str, author: str, category: str, bg_color: tuple, filename: str) -> str:
    path = os.path.join(COVERS_DIR, filename)
    width, height = 600, 850
    img = Image.new("RGB", (width, height), bg_color)
    draw = ImageDraw.Draw(img)
    
    # Decorative elements
    draw.rectangle([20, 20, width - 20, height - 20], outline=(255, 255, 255, 180), width=3)
    draw.rectangle([40, 40, width - 40, 140], fill=(0, 0, 0, 80))
    
    # Category badge
    draw.rectangle([60, 60, 260, 100], fill=(255, 255, 255))
    draw.text((80, 72), category.upper(), fill=(0, 0, 0))
    
    # Title & Author
    draw.text((60, 220), title[:25], fill=(255, 255, 255))
    if len(title) > 25:
        draw.text((60, 260), title[25:], fill=(255, 255, 255))
        
    draw.text((60, 360), "DIGITAL EDITION • INSTANT ACCESS", fill=(220, 220, 240))
    
    # Footer
    draw.line([(60, 720), (width - 60, 720)], fill=(255, 255, 255), width=2)
    draw.text((60, 750), f"Author: {author}", fill=(255, 255, 255))
    draw.text((60, 780), "EBookVault Exclusive Edition", fill=(200, 200, 220))
    
    img.save(path, "JPEG", quality=90)
    return f"/uploads/covers/{filename}"

def generate_sample_pdf(title: str, author: str, filename: str) -> str:
    path = os.path.join(EBOOKS_DIR, filename)
    doc = SimpleDocTemplate(path, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'EbookTitle',
        parent=styles['Heading1'],
        fontSize=26,
        leading=32,
        textColor=colors.HexColor('#1e1b4b'),
        spaceAfter=15
    )
    author_style = ParagraphStyle(
        'EbookAuthor',
        parent=styles['Normal'],
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#4338ca'),
        spaceAfter=25
    )
    heading2_style = ParagraphStyle(
        'EbookHeading2',
        parent=styles['Heading2'],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1f2937'),
        spaceBefore=20,
        spaceAfter=10
    )
    body_style = ParagraphStyle(
        'EbookBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        textColor=colors.HexColor('#374151'),
        spaceAfter=12
    )
    
    story = []
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"By {author}", author_style))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("Chapter 1: The Modern Foundation", heading2_style))
    story.append(Paragraph(
        "Welcome to this comprehensive guide. In the fast-moving landscape of digital technology, "
        "mastery begins with core principles: clarity of purpose, systematic execution, and continuous optimization.",
        body_style
    ))
    story.append(Paragraph(
        "Whether you are building high-scale automated workflows or crafting modern software products, "
        "the secret lies in modular architecture. When every component has a single, well-defined responsibility, "
        "your entire system becomes resilient, maintainable, and adaptable.",
        body_style
    ))
    
    story.append(Paragraph("Chapter 2: Step-by-Step Practical Implementation", heading2_style))
    story.append(Paragraph(
        "1. Define the end-state architecture before writing the first line of code.<br/>"
        "2. Automate repetitive delivery pipelines to eliminate human error.<br/>"
        "3. Provide instant multi-channel feedback loops to delight your end users.<br/>"
        "4. Continuously measure performance and user satisfaction.",
        body_style
    ))
    story.append(Paragraph(
        "Thank you for reading! Keep this document safe as reference for your future implementations.",
        body_style
    ))
    
    doc.build(story)
    return path

def generate_sample_docx(title: str, author: str, filename: str) -> str:
    path = os.path.join(EBOOKS_DIR, filename)
    doc = Document()
    
    # Title
    p_title = doc.add_heading(title, level=0)
    p_title.runs[0].font.color.rgb = RGBColor(30, 27, 75)
    
    # Author
    p_author = doc.add_paragraph(f"By {author} | Official Digital Edition")
    p_author.runs[0].font.italic = True
    p_author.runs[0].font.color.rgb = RGBColor(79, 70, 229)
    
    doc.add_heading("Table of Contents", level=2)
    doc.add_paragraph("1. Introduction & Strategy Overview\n2. Key Execution Frameworks\n3. Scaling Your Operations\n4. Case Studies & Blueprints")
    
    doc.add_heading("Chapter 1: The Strategic Blueprint", level=1)
    doc.add_paragraph(
        "Every successful digital business or project relies on three core pillars: "
        "A compelling product, friction-free distribution, and immediate automated delivery. "
        "When buyers purchase your product, instant gratification builds trust and loyalty."
    )
    
    doc.add_heading("Chapter 2: The Multi-Channel Advantage", level=1)
    doc.add_paragraph(
        "Delivering your digital products through both Email and WhatsApp ensures that no customer "
        "ever loses access. Inboxes can get cluttered, but WhatsApp messages offer instant visibility and one-tap re-downloads."
    )
    
    doc.save(path)
    return path

async def seed():
    from database import init_db
    await init_db()
    
    sample_books = [
        {
            "title": "Mastering Python & AI Automation 2026",
            "slug": "mastering-python-ai-automation",
            "author": "Dr. Sarah Jenkins",
            "description": "The definitive, hands-on guide to building scalable autonomous systems, AI workflows, and background job queues with modern Python.",
            "price": 29.99,
            "sale_price": 19.99,
            "category": "Technology & Coding",
            "bg_color": (30, 41, 59),
            "cover_file": "python-ai-cover.jpg",
            "format": "pdf",
            "doc_file": "mastering-python-ai-automation.pdf",
            "sample_text": "Chapter 1: The Modern Foundation\nWelcome to this comprehensive guide. In the fast-moving landscape of digital technology, mastery begins with core principles...",
            "is_featured": 1
        },
        {
            "title": "The Modern Solopreneur Blueprint",
            "slug": "modern-solopreneur-blueprint",
            "author": "Marcus Vance",
            "description": "Step-by-step framework to launch, monetize, and scale digital products and automated consulting to $10,000/month with zero employees.",
            "price": 39.00,
            "sale_price": 24.50,
            "category": "Business & Startups",
            "bg_color": (15, 23, 42),
            "cover_file": "solopreneur-blueprint.jpg",
            "format": "docx",
            "doc_file": "modern-solopreneur-blueprint.docx",
            "sample_text": "Chapter 1: The Strategic Blueprint\nEvery successful digital business relies on three core pillars: A compelling product, friction-free distribution, and immediate automated delivery...",
            "is_featured": 1
        },
        {
            "title": "Zero to Financial Freedom: Modern Investing",
            "slug": "zero-to-financial-freedom",
            "author": "Elena Rostova",
            "description": "Master index investing, real estate cash flow, asymmetric risk strategies, and tax-efficient wealth compounding for high earners.",
            "price": 34.99,
            "sale_price": 22.00,
            "category": "Personal Finance",
            "bg_color": (19, 78, 74),
            "cover_file": "financial-freedom.jpg",
            "format": "pdf",
            "doc_file": "financial-freedom-investing.pdf",
            "sample_text": "Understanding Cash Flow & Compound Wealth: The fundamental math behind financial independence is simpler than Wall Street wants you to believe...",
            "is_featured": 1
        },
        {
            "title": "Freelancer Mastery: High-Ticket Acquisition",
            "slug": "freelancer-mastery-high-ticket",
            "author": "Alex Rivera",
            "description": "Stop competing on low-rate platforms. Learn inbound outreach, executive positioning, and value-based closing for 5-figure projects.",
            "price": 27.00,
            "sale_price": 17.50,
            "category": "Career & Freelancing",
            "bg_color": (88, 28, 135),
            "cover_file": "freelancer-mastery.jpg",
            "format": "docx",
            "doc_file": "freelancer-mastery-guide.docx",
            "sample_text": "The High-Ticket Paradigm Shift: When you sell commodities, price is a race to the bottom. When you sell business outcomes, price becomes an investment...",
            "is_featured": 0
        }
    ]
    
    async with aiosqlite.connect(DB_PATH) as db:
        for book in sample_books:
            # Generate cover
            cover_url = generate_cover_image(
                book["title"],
                book["author"],
                book["category"],
                book["bg_color"],
                book["cover_file"]
            )
            
            # Generate file
            if book["format"] == "pdf":
                file_path = generate_sample_pdf(book["title"], book["author"], book["doc_file"])
            else:
                file_path = generate_sample_docx(book["title"], book["author"], book["doc_file"])
                
            file_size = os.path.getsize(file_path)
            
            await db.execute("""
                INSERT INTO ebooks (
                    title, slug, author, description, price, sale_price, category,
                    cover_image, file_path, file_name, file_format, file_size_bytes,
                    sample_text, is_featured, is_active
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(slug) DO UPDATE SET
                    price = excluded.price,
                    sale_price = excluded.sale_price,
                    cover_image = excluded.cover_image,
                    file_path = excluded.file_path,
                    file_format = excluded.file_format,
                    file_size_bytes = excluded.file_size_bytes
            """, (
                book["title"],
                book["slug"],
                book["author"],
                book["description"],
                book["price"],
                book["sale_price"],
                book["category"],
                cover_url,
                file_path,
                book["doc_file"],
                book["format"],
                file_size,
                book["sample_text"],
                book["is_featured"],
                1
            ))
            
        await db.commit()
    print("Database and starter sample ebooks seeded successfully!")

if __name__ == "__main__":
    asyncio.run(seed())
